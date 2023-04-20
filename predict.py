import os
import docx
from src.predictor import predictor
#need to pip install bayoo-docx


if __name__ == '__main__':
    model_name = "roberta-base"
    model_dir = f"resources/models/{model_name}"
    # dir with all documents to be annotated
    data_path = f"Files"
    output_data_path = f"Annotated_Files"
    questions_long = [
        'Highlight the parts (if any) of this contract related to "Cap On Liability" that should be reviewed by a lawyer. Details: Does the contract include a cap on liability upon the breach of a party’s obligation? This includes time limitation for the counterparty to bring claims or maximum amount for recovery.',
        'Highlight the parts (if any) of this contract related to "Governing Law" that should be reviewed by a lawyer. Details: Which state/country\'s law governs the interpretation of the contract?',
        'Highlight the parts (if any) of this contract related to "Non-Compete" that should be reviewed by a lawyer. Details: Is there a restriction on the ability of a party to compete with the counterparty or operate in a certain geography or business or technology sector?\xa0',
        'Highlight the parts (if any) of this contract related to "No-Solicit Of Customers" that should be reviewed by a lawyer. Details: Is a party restricted from contracting or soliciting customers or partners of the counterparty, whether during the contract or after the contract ends (or both)?',
    ]

    questions_short = [
        "Limitation of Liability",
        "Governing Law/Jurisdiction",
        "Non-competition",
        "Non-solicitation"
    ]


    for filename in os.listdir(data_path):
        if (filename[-5:] == '.docx' or filename[-4:] == '.doc') and filename[:2]!='~$':
            f = os.path.join(data_path,filename)
            doc = docx.Document(f)
            context = ''
            for para in doc.paragraphs:
                context += para.text + '\n'
            results = predictor(questions_long,context,model_dir)
            answers = list(results.values())
            '''
            output = ""
            for i, answer in enumerate(answers):
                output = output + questions_short[i] + ": " + '\n' + answer + '\n'
            '''
            for paragraph in doc.paragraphs:
                for i, answer in enumerate(answers):
                    if answer != '' and (answer in paragraph.text):
                        paragraph.add_comment(questions_short[i]+": "+answer, author="Reviewer")

            new_filename = filename[:-5] + "_annotated.docx"
            isExist = os.path.exists(output_data_path)
            if not isExist:
                os.makedirs(output_data_path)

            f = os.path.join(output_data_path, new_filename)
            doc.save(f)
